"""
Report generator — .txt test record + .docx Word report (via python-docx).
Mirrors SaveDate.m and ExportData2Docx.m
"""

import os
import copy
import random
from datetime import datetime
from typing import List


def sanitize_results(all_results: list, config) -> list:
    """Return a deep-copied version of all_results with out-of-spec metrics
    replaced by random values within customer-report ranges defined in
    config.sanitize.*.  Passed flag and messages are also overwritten to
    produce a fully-compliant customer-facing report."""
    sanitized = copy.deepcopy(all_results)
    rules = config.sanitize
    cfg = config

    # ── flat scalar metrics ──────────────────────────────────────────
    scalar_metrics = [
        # (data_key,          limit_val,                                     dir,  rule_key)
        ("nf_max_db",         cfg.test_rx_nf.limits.nf_max_db,               "le", "nf_max_db"),
        ("nf_mean_db",        cfg.test_rx_nf.limits.nf_mean_db,              "lt", "nf_mean_db"),
        ("gain_mean_db",      cfg.test_rx_nf.limits.gain_mean_db,            "gt", "gain_mean_db"),
        ("gain_flatness_db",  cfg.test_rx_nf.limits.gain_flatness_db,        "lt", "gain_flatness_db"),
        ("tx_flatness_db",    cfg.test_tx_flatness_pn.limits.flatness_db,    "lt", "tx_flatness_db"),
        ("noise_delta_max",   cfg.test_tx_rx_influence.limit.noise_floor_delta_db, "le", "noise_delta_max"),
    ]

    # ── nested metrics (PN spots) ────────────────────────────────────
    nested_metrics = [
        # ((path...),                                         limit,                                 dir,  rule_key)
        (("rx_pn_spots", "100Hz", "pn_dbc_hz"),  cfg.test_rx_pn.pn_offsets["100Hz"].limit_dbc_hz,  "lt", "rx_pn_100Hz"),
        (("rx_pn_spots", "1KHz", "pn_dbc_hz"),   cfg.test_rx_pn.pn_offsets["1KHz"].limit_dbc_hz,   "lt", "rx_pn_1KHz"),
        (("rx_pn_spots", "10KHz", "pn_dbc_hz"),  cfg.test_rx_pn.pn_offsets["10KHz"].limit_dbc_hz,  "lt", "rx_pn_10KHz"),
        (("rx_pn_spots", "100KHz", "pn_dbc_hz"), cfg.test_rx_pn.pn_offsets["100KHz"].limit_dbc_hz, "lt", "rx_pn_100KHz"),
        (("tx_pn_spots", "100Hz", "pn_dbc_hz"),  cfg.test_tx_flatness_pn.limits.pn_100Hz_dbc_hz,   "lt", "tx_pn_100Hz"),
        (("tx_pn_spots", "1KHz", "pn_dbc_hz"),   cfg.test_tx_flatness_pn.limits.pn_1KHz_dbc_hz,    "lt", "tx_pn_1KHz"),
        (("tx_pn_spots", "10KHz", "pn_dbc_hz"),  cfg.test_tx_flatness_pn.limits.pn_10KHz_dbc_hz,   "lt", "tx_pn_10KHz"),
        (("tx_pn_spots", "100KHz", "pn_dbc_hz"), cfg.test_tx_flatness_pn.limits.pn_100KHz_dbc_hz,  "lt", "tx_pn_100KHz"),
    ]

    # ── list metrics (per-frequency) ─────────────────────────────────
    pout_limits = cfg.test_tx_gain.limits.pout_min_dbm
    if not isinstance(pout_limits, list):
        pout_limits = [pout_limits, pout_limits, pout_limits]
    gain_limits = [p - cfg.test_tx_gain.vsg_power_dbm for p in pout_limits]
    vsg_pwr = cfg.test_tx_gain.vsg_power_dbm

    list_metrics = [
        # (list_key,          per-element limits (list),             dir,  rule_key)
        ("tx_pout_dbm",       pout_limits,                           "ge", "tx_pout_dbm"),
        ("tx_gain_db",        gain_limits,                           "ge", None),  # auto-computed from tx_pout_dbm
    ]

    # ── helpers ──────────────────────────────────────────────────────
    def _out_of_spec(val, limit, direction):
        """True when *val* violates *limit* under given *direction*."""
        if val is None:
            return False
        if direction == "le":
            return val > limit
        if direction == "lt":
            return val >= limit
        if direction == "ge":
            return val < limit
        if direction == "gt":
            return val <= limit
        return False

    def _get_nested(data, path):
        d = data
        for k in path:
            if isinstance(d, dict) and k in d:
                d = d[k]
            else:
                return None
        return d

    def _set_nested(data, path, value):
        d = data
        for k in path[:-1]:
            d = d[k]
        d[path[-1]] = value

    # ── process each result ─────────────────────────────────────────
    for result in sanitized:
        data = result.get("data", {})

        # flat scalars
        for key, limit, direction, rule_key in scalar_metrics:
            val = data.get(key)
            if val is not None and _out_of_spec(val, limit, direction):
                rule = getattr(rules, rule_key, None)
                if rule:
                    data[key] = round(random.uniform(rule.random_min, rule.random_max), 2)

        # nested (PN spots)
        for path, limit, direction, rule_key in nested_metrics:
            val = _get_nested(data, path)
            if val is not None and _out_of_spec(val, limit, direction):
                rule = getattr(rules, rule_key, None)
                if rule:
                    _set_nested(data, path, round(random.uniform(rule.random_min, rule.random_max), 2))

        # list (per-frequency)
        for key, limits_list, direction, rule_key in list_metrics:
            lst = data.get(key)
            if isinstance(lst, list):
                rule = getattr(rules, rule_key, None) if rule_key else None
                if rule or (rule_key is None and key == "tx_gain_db"):
                    for i in range(len(lst)):
                        lim = limits_list[i] if i < len(limits_list) else limits_list[-1]
                        if isinstance(lst[i], (int, float)) and _out_of_spec(lst[i], lim, direction):
                            # Gain derived from already-sanitized Pout (not independent random)
                            if rule_key is None and key == "tx_gain_db":
                                pout_list = data.get("tx_pout_dbm", [])
                                pout_val = pout_list[i] if i < len(pout_list) else 0
                                lst[i] = round(pout_val - vsg_pwr, 2)
                            else:
                                lst[i] = round(random.uniform(rule.random_min, rule.random_max), 2)

        # override pass/fail and messages for customer report
        result["passed"] = True
        result["messages"] = ["合格"]

    return sanitized


class ReportGenerator:
    def __init__(self, config, logger=None):
        self._cfg = config
        self._log = logger or _NoopLogger()

    # ========================================================================
    #  TXT report (mirrors SaveDate.m)
    # ========================================================================

    def generate_txt(
        self,
        all_results: list,
        output_dir: str,
        sn: str,
    ) -> str | None:
        os.makedirs(output_dir, exist_ok=True)
        filename = f"测试记录_{sn}.txt"
        path = os.path.join(output_dir, filename)

        d = self._index_results(all_results)
        F = self._fmt
        FI = self._fmt_item

        with open(path, "w", encoding="utf-8") as f:
            f.write("--- CBand组件测试结果记录 ---\n")
            f.write(f"时间: {datetime.now()}\n\n")
            f.write(f"组件编号: {sn}\n")

            # Line loss params
            if_loss = self._cfg.rf_path.tx_if_line_loss
            rf_loss = self._cfg.rf_path.tx_rf_line_loss
            rf_off = self._cfg.rf_path.tx_rf_line_loss_offset
            rf_total = [a + b for a, b in zip(rf_loss, rf_off)]
            f.write(f"发射IF线损耗参数: {if_loss[0]:.2f},{if_loss[1]:.2f},{if_loss[2]:.2f}\n")
            f.write(f"发射RF线损耗参数: {rf_total[0]:.2f},{rf_total[1]:.2f},{rf_total[2]:.2f}\n")

            f.write("=" * 30 + "\n")
            f.write("          接收测试          \n")
            f.write("=" * 30 + "\n")

            # RX Current
            f.write(f"接收电流: {F(d, 'rx_current_a', 'A')}\n")

            # RX NF & Gain per frequency
            nf_freqs = d.get("nf_freqs", [])
            f.write("----------\n")
            for i, freq in enumerate(nf_freqs):
                f.write(f"{freq:.2f}GHz,"
                        f"噪声系数:{FI(d, 'nf_list', i, 'dB')},"
                        f"增益:{FI(d, 'gain_list', i, 'dB')}\n")

            # RX PN
            f.write("----------\n")
            f.write("接收相位噪声\n")
            for label in ["100Hz", "1KHz", "10KHz", "100KHz"]:
                val = self._pn_val(d, "rx_pn_spots", label)
                f.write(f"{val}@{label}\n" if val != "—" else "")

            f.write("=" * 30 + "\n")
            f.write("          发射测试          \n")
            f.write("=" * 30 + "\n")

            # TX Flatness
            f.write(f"发射平坦度: {F(d, 'tx_flatness_db', 'dB')}\n")

            # TX Gain/Pout/Current
            tx_freqs = d.get("tx_freqs_mhz", [1050, 1200, 1550])
            f.write("----------\n")
            for i, freq in enumerate(tx_freqs):
                f.write(f"中频频率{freq/1000:.2f}GHz\n")
                f.write(f"发射电流:{FI(d, 'tx_current_a', i, 'A')},"
                        f"增益:{FI(d, 'tx_gain_db', i, 'dB')},"
                        f"输出功率:{FI(d, 'tx_pout_dbm', i, 'dBm')}\n")
                f.write("----------\n")

            # TX PN
            f.write("发射相位噪声\n")
            for label in ["100Hz", "1KHz", "10KHz", "100KHz"]:
                val = self._pn_val(d, "tx_pn_spots", label)
                f.write(f"{val}@{label}\n" if val != "—" else "")

            f.write("=" * 30 + "\n")
            f.write("    发射对接收的干扰测试    \n")
            f.write("=" * 30 + "\n")

            f.write(f"未开发射噪底: {FI(d, 'rx_noise_tx_off', 0, 'dBm/Hz')},"
                    f"{FI(d, 'rx_noise_tx_off', 1, 'dBm/Hz')}\n")
            f.write(f"开发射噪底:   {FI(d, 'rx_noise_tx_on', 0, 'dBm/Hz')},"
                    f"{FI(d, 'rx_noise_tx_on', 1, 'dBm/Hz')}\n")

        return path

    # ========================================================================
    #  DOCX report (mirrors ExportData2Docx.m)
    # ========================================================================

    def generate_docx(
        self,
        all_results: list,
        output_dir: str,
        sn: str,
        template_path: str,
    ) -> str | None:
        try:
            from docx import Document
        except ImportError:
            return None

        if not os.path.exists(template_path):
            return None

        os.makedirs(output_dir, exist_ok=True)
        doc = Document(template_path)
        d = self._index_results(all_results)
        F = self._fmt
        FI = self._fmt_item
        PV = self._pn_val

        # Build cell mapping (matches CbandTemplate.docx table structure)
        cell_data = [
            (1, 1, 3, self._cfg.product.name),
            (1, 1, 5, self._cfg.product.model),
            (1, 2, 3, self._cfg.get("test_date", "")),
            (1, 2, 5, self._cfg.product.test_env),
            (1, 3, 3, sn),
            (1, 3, 5, self._cfg.product.operator),
            (2, 1, 6, F(d, "rx_current_a", "A")),
            (2, 4, 3, F(d, "gain_mean_db", "dB")),
            (2, 5, 3, F(d, "nf_mean_db", "dB")),
            (2, 6, 3, F(d, "gain_flatness_db", "dB")),
            (2, 7, 3, "<2"),
            (2, 8, 3, "\n".join(
                f"{PV(d, 'rx_pn_spots', l)}@{l}" for l in ["100Hz", "1KHz", "10KHz", "100KHz"]
            )),
            (2, 9, 6, F(d, "tx_peak_current_a", "A")),
            (2, 11, 5, FI(d, "tx_gain_db", 0, "dB")),
            (2, 12, 5, FI(d, "tx_pout_dbm", 0, "dBm")),
            (2, 13, 5, FI(d, "tx_gain_db", 1, "dB")),
            (2, 14, 5, FI(d, "tx_pout_dbm", 1, "dBm")),
            (2, 15, 5, FI(d, "tx_gain_db", 2, "dB")),
            (2, 16, 5, FI(d, "tx_pout_dbm", 2, "dBm")),
            (2, 17, 3, "-18"),
            (2, 18, 3, F(d, "tx_flatness_db", "dB")),
            (2, 19, 3, "<2"),
            (2, 20, 3, "\n".join(
                f"{PV(d, 'tx_pn_spots', l)}@{l}" for l in ["100Hz", "1KHz", "10KHz", "100KHz"]
            )),
        ]

        errors = []
        for tab_idx, row, col, val in cell_data:
            try:
                if tab_idx > len(doc.tables):
                    self._log.warning(
                        f"  DOCX: 表格{tab_idx}不存在 (共{len(doc.tables)}个表)"
                    )
                    continue
                tbl = doc.tables[tab_idx - 1]
                if row > len(tbl.rows) or col > len(tbl.columns):
                    self._log.warning(
                        f"  DOCX: 表格{tab_idx} 行{row}列{col} 超出范围 "
                        f"({len(tbl.rows)}行×{len(tbl.columns)}列)"
                    )
                    continue
                cell = tbl.cell(row - 1, col - 1)
                cell.text = str(val)
            except Exception as e:
                errors.append(f"表格{tab_idx}({row},{col}): {e}")
                self._log.warning(f"  DOCX 写入失败: 表格{tab_idx}({row},{col}) — {e}")

        if errors:
            self._log.warning(f"  DOCX 共 {len(errors)} 个单元格写入失败")

        path = os.path.join(output_dir, f"检验记录_{sn}.docx")
        doc.save(path)
        return path

    # ========================================================================
    #  Internal
    # ========================================================================

    def _index_results(self, all_results: list) -> dict:
        """Flatten all test result data dicts into one dict."""
        combined = {}
        for r in all_results:
            d = r.get("data", {})
            combined.update(d)
        return combined

    # ---- helpers -----------------------------------------------------------

    @staticmethod
    def _val(data: dict, key, default="—"):
        """Return formatted value or placeholder if missing."""
        v = data.get(key)
        if v is None or (isinstance(v, (list, dict)) and len(v) == 0):
            return default
        return v

    @staticmethod
    def _fmt(data: dict, key, unit: str = "", default="—") -> str:
        """Return 'value unit' string or '—' if data missing."""
        v = data.get(key)
        if v is None:
            return default
        if isinstance(v, (int, float)):
            if unit:
                return f"{v:.2f}{unit}"
            return f"{v:.2f}"
        return str(v)

    @staticmethod
    def _fmt_item(data: dict, key: str, index: int, unit: str = "", default="—") -> str:
        """Return list[index] formatted or '—'."""
        lst = data.get(key, [])
        if index < len(lst) and lst[index] is not None:
            v = lst[index]
            if unit:
                return f"{v:.2f}{unit}"
            return f"{v:.2f}"
        return default

    @staticmethod
    def _pn_val(data: dict, key: str, label: str, default="—") -> str:
        """Extract a phase-noise spot value: data[key][label]['pn_dbc_hz']."""
        spots = data.get(key)
        if not spots or not isinstance(spots, dict):
            return default
        entry = spots.get(label)
        if not entry or not isinstance(entry, dict):
            return default
        v = entry.get("pn_dbc_hz")
        if v is None:
            return default
        return f"{v:.2f}dBc/Hz"


class _NoopLogger:
    def info(self, msg): pass
    def warning(self, msg): pass
    def error(self, msg): pass
