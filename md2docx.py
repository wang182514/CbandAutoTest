"""Convert 测试手册.md to 测试手册.docx using python-docx."""
import re
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# ── style setup ──
style = doc.styles['Normal']
style.font.name = 'Microsoft YaHei'
style.font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(4)

# ── parse and build ──
with open('测试手册.md', 'r', encoding='utf-8') as f:
    lines = f.readlines()

i = 0
current_table = []
in_code = False

def add_heading(text, level):
    h = doc.add_heading(text.strip(), level=level)
    return h

def add_para(text):
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)  # bold
    text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)  # links
    text = re.sub(r'`(.+?)`', r'\1', text)  # inline code
    doc.add_paragraph(text.strip())

def flush_table():
    global current_table
    if not current_table:
        return
    rows_data = []
    for row in current_table:
        cells = [c.strip() for c in row.split('|')[1:-1]]
        if all(c.replace('-','').replace(':','').strip() == '' for c in cells):
            continue
        rows_data.append(cells)
    if len(rows_data) < 2:
        current_table = []
        return
    ncols = len(rows_data[0])
    tbl = doc.add_table(rows=len(rows_data), cols=ncols)
    tbl.style = 'Light Grid Accent 1'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ri, row_cells in enumerate(rows_data):
        for ci, cell_text in enumerate(row_cells):
            if ci < ncols:
                cell = tbl.cell(ri, ci)
                cell.text = cell_text
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        run.font.size = Pt(9)
    current_table = []

while i < len(lines):
    line = lines[i]

    # code block
    if line.strip().startswith('```'):
        in_code = not in_code
        i += 1
        continue
    if in_code:
        add_para(line.rstrip())
        i += 1
        continue

    # table — accumulate rows
    if line.strip().startswith('|'):
        current_table.append(line.rstrip())
        # peek next line
        if i + 1 < len(lines) and lines[i+1].strip().startswith('|'):
            i += 1
            continue
        else:
            flush_table()
            i += 1
            continue

    # heading
    if line.startswith('#'):
        level = len(line) - len(line.lstrip('#'))
        heading_text = line.lstrip('#').strip()
        # skip version line
        if heading_text.startswith('>'):
            add_para(heading_text.lstrip('>').strip())
        else:
            add_heading(heading_text, min(level, 3))
        i += 1
        continue

    # blockquote
    if line.strip().startswith('>'):
        add_para(line.strip().lstrip('>').strip())
        i += 1
        continue

    # horizontal rule
    if line.strip() == '---':
        flush_table()
        doc.add_paragraph()
        i += 1
        continue

    # blank line
    if line.strip() == '':
        i += 1
        continue

    # normal paragraph
    flush_table()
    add_para(line.rstrip())
    i += 1

# flush any remaining table
flush_table()

# ── save ──
doc.save('测试手册.docx')
print('Done → 测试手册.docx')
